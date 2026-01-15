#!/usr/bin/env python3
"""
Resume Generator
Generates a professional resume DOCX from JSON data using python-docx.
"""
import json
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def load_resume_data(filepath="resume_data.json"):
    """Load resume data from JSON file."""
    try:
        with open(filepath, "r") as f:
            return json.load(f)
    except FileNotFoundError:
        print(f"Error: {filepath} not found.")
        return {}
    except json.JSONDecodeError:
        print(f"Error: Failed to decode JSON from {filepath}.")
        return {}

def hex_to_rgb(hex_color):
    """Convert hex string to RGB tuple."""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def add_bottom_border(paragraph):
    """Add a bottom border to a paragraph."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)

def setup_document():
    """Initialize document and section margins."""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    return doc

def setup_styles(doc):
    """Configure Normal style for consistent formatting."""
    styles = doc.styles
    
    # Normal Style
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(2)

def add_contact_info(doc, contact_data):
    """Add contact info as a single line with separators."""
    if not contact_data:
        return
    
    # Check if full address should be shown
    show_full_address = os.environ.get("RESUME_FULL_ADDRESS", "").lower() in ("true", "1", "yes")
        
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    parts = []
    # Extract details manually or join if it represents lines
    if "details" in contact_data:
        # If details is a string, just use it
        parts.append(contact_data["details"])
    
    # Address handling based on env var
    if show_full_address:
        # Full address mode
        addr = []
        if "address_line1" in contact_data: addr.append(contact_data["address_line1"])
        if "address_line2" in contact_data: addr.append(contact_data["address_line2"])
        if addr:
            parts.insert(0, ", ".join(addr))
    else:
        # Redacted mode - only show state
        if "address_state" in contact_data:
            parts.insert(0, contact_data["address_state"])
        
    # Join with separators
    full_text = " | ".join(parts)
    run = p.add_run(full_text)
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(12)

def add_section_header(doc, title):
    """Add a styled section header with border."""
    p = doc.add_paragraph()
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(12)
    # Green accent color #4d9a47
    run.font.color.rgb = RGBColor(*hex_to_rgb("4d9a47"))
    
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    add_bottom_border(p)
    return p

def add_job_header(doc, role, company, dates):
    """Add a job header with role, company, and right-aligned date."""
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Inches(7.3), WD_TAB_ALIGNMENT.RIGHT)
    
    # Role (Bold)
    r1 = p.add_run(role)
    r1.bold = True
    r1.font.size = Pt(11)
    
    p.add_run(", ")
    
    # Company (Normal or Italic)
    r2 = p.add_run(company)
    r2.italic = False # Can be true if preferred
    
    # Push date to right
    p.add_run("\t")
    r3 = p.add_run(dates)
    r3.italic = True
    
    p.paragraph_format.space_after = Pt(2)
    return p

def add_bullet(doc, text):
    """Add a bullet point with custom indent."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.paragraph_format.space_after = Pt(1)
    
    run_bullet = p.add_run("• ")
    run_bullet.font.color.rgb = RGBColor(*hex_to_rgb("4d9a47")) # Green bullet
    
    p.add_run(text)
    return p

def build_resume(data, output_path="public/Metehan_Ozten_Resume.docx"):
    """Build the resume DOCX from data."""
    if not data:
        return

    doc = setup_document()
    setup_styles(doc)

    # --- Name ---
    if "name" in data:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(data["name"])
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(*hex_to_rgb("2d5a27")) # Darker green for name
        p.paragraph_format.space_after = Pt(4)
    
    # --- Contact ---
    add_contact_info(doc, data.get("contact", {}))

    # --- Summary ---
    if "summary" in data and data["summary"]:
        add_section_header(doc, "Summary")
        p = doc.add_paragraph(data["summary"])
        p.paragraph_format.space_after = Pt(2)

    # --- Education ---
    if "education" in data:
        add_section_header(doc, "Education")
        for edu in data["education"]:
            # Institution
            if "institution" in edu:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                r = p.add_run(edu["institution"])
                r.bold = True
                r.font.size = Pt(11)

            # Items (Degree, etc)
            if "items" in edu:
                for item in edu["items"]:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(0)
                    # Support left/right split if needed, otherwise just text
                    left = item.get("left", "")
                    right = item.get("right")
                    
                    if right:
                        p.paragraph_format.tab_stops.add_tab_stop(Inches(7.3), WD_TAB_ALIGNMENT.RIGHT)
                        p.add_run(left + "\t" + right)
                    else:
                        p.add_run(left)
            doc.add_paragraph("") # Spacing after each edu entry

    # --- Experience ---
    if "experience" in data:
        add_section_header(doc, "Experience")
        for job in data["experience"]:
            role = job.get("role", "")
            company = job.get("company", "")
            dates = job.get("dates", "")
            
            add_job_header(doc, role, company, dates)
            
            for b in job.get("bullets", []):
                add_bullet(doc, b)
            
            # Small spacer between jobs
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)

    # --- Projects ---
    if "projects" in data:
        add_section_header(doc, "Projects")
        for project in data["projects"]:
            title = project.get("title", "")
            url = project.get("url", "")
            
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            
            r_title = p.add_run(title)
            r_title.bold = True
            
            if url:
                r_sep = p.add_run(" | ")
                r_sep.font.color.rgb = RGBColor(180, 180, 180)
                r_url = p.add_run(url)
                r_url.italic = True
                r_url.font.color.rgb = RGBColor(*hex_to_rgb("0077b5")) # Link blue
                
            for b in project.get("bullets", []):
                add_bullet(doc, b)
            
            # In Progress
            in_progress = project.get("in_progress", [])
            if in_progress:
                p_sub = doc.add_paragraph()
                p_sub.paragraph_format.left_indent = Inches(0.15)
                r_sub = p_sub.add_run("In Progress:")
                r_sub.italic = True
                r_sub.font.size = Pt(10)
                
                for ip in in_progress:
                    add_bullet(doc, ip)
            
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)

    # --- Skills ---
    if "skills" in data:
        add_section_header(doc, "Technical Skills")
        skills = data["skills"]
        if isinstance(skills, dict):
            for k, v in skills.items():
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                
                cat = p.add_run(f"{k}: ")
                cat.bold = True
                
                content = ", ".join(v) if isinstance(v, list) else v
                p.add_run(content)

    # --- Distinctions & Leadership (Grouped if desired, or separate) ---
    if "distinctions" in data:
        add_section_header(doc, "Distinctions")
        for d in data["distinctions"]:
            add_bullet(doc, d)

    if "leadership" in data:
        add_section_header(doc, "Leadership")
        for l in data["leadership"]:
            add_bullet(doc, l)

    # --- Publications ---
    if "publications" in data:
        add_section_header(doc, "Publications")
        for line in data["publications"]:
            p = doc.add_paragraph()
            p.add_run(line)

    # Save
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"✅ Resume generated successfully: {output_path}")

def main():
    """Main entry point."""
    script_dir = os.path.dirname(os.path.abspath(__file__))
    os.chdir(script_dir)
    
    data = load_resume_data()
    build_resume(data)

if __name__ == "__main__":
    main()
